#!/usr/bin/env python3
"""
Responsibility:
Extracts and sanitizes text from Word documents (DOCX natively; legacy DOC via external tools) and converts the result into fixed-size attachment chunk tasks.

Used by:
* rag/chunk_att.py

Pipelines:
- read_docx -> sanitize_text -> join_segments -> chunk_fixed -> build_tasks
- run_tool -> sanitize_text -> join_segments -> chunk_fixed -> build_tasks

Invariants:
- Only filenames with extensions in `WORD_EXTS` produce tasks.
- Legacy `.doc` extraction attempts `antiword`, then `catdoc`, then `soffice/libreoffice` with timeouts.
- Returned task tuples contain 1-based `seq` values within a single document.

Out of scope:
- Email attachment iteration and routing (handled by `chunk_att`).
- PDF/Excel extraction (handled by `chunk_pdf` / `chunk_xls`).
"""

import io
import logging
import shutil
import subprocess
import tempfile
from pathlib import Path
from typing import List, Tuple

import docx
from chunk_att import build_attachment_tasks, join_nonempty_segments
from helper.helper_sanitize import sanitize_text

logger = logging.getLogger(__name__)

WORD_EXTS = {".doc", ".docx"}


def _extract_text_from_docx(data: bytes) -> dict[int, str]:
    """
    Purpose:
    Extract and sanitize text from DOCX bytes.

    Inputs:
    - data: Raw DOCX file bytes.

    Outputs:
    - Mapping of 1-based paragraph index to sanitized paragraph text; may fall back to a single entry built from tables.

    Side effects:
    - Emits log messages about extraction outcomes.

    Failure modes:
    - Returns `{}` when parsing fails or no extractable text is found.
    """

    try:
        doc = docx.Document(io.BytesIO(data))

        # 1) 先取段落
        paragraphs: dict[int, str] = {
            i: sanitize_text(p.text.strip())
            for i, p in enumerate(doc.paragraphs, 1)
            if p.text.strip()
        }

        # 2) 如无段落，做表格兜底（行内用 TAB，表间换行）
        if not paragraphs and getattr(doc, "tables", None):
            rows: list[str] = []
            for t in doc.tables:
                for row in t.rows:
                    rows.append("\t".join(cell.text.strip() for cell in row.cells))
            text = "\n".join(r for r in rows if r.strip())
            if text:
                paragraphs = {1: sanitize_text(text)}
            else:
                logger.warning("Empty or unreadable .docx file (no paragraphs/tables)")
        elif not paragraphs:
            logger.warning("Empty or unreadable .docx file")

        logger.info("Word extraction complete: %d paragraphs", len(paragraphs))
        return paragraphs
    except Exception as exc:  # noqa: BLE001
        logger.error("Failed to process .docx document: %s", exc)
        return {}


def extract_text_from_doc(data: bytes, *, timeout: int = 15) -> dict[int, str]:
    """
    Purpose:
    Extract and sanitize text from legacy `.doc` bytes via external tools with a timeout.

    Inputs:
    - data: Raw `.doc` file bytes.
    - timeout: Maximum seconds to allow any single subprocess attempt to run.

    Outputs:
    - Mapping of 1-based paragraph index to sanitized paragraph text.

    Side effects:
    - Runs external subprocesses when available (`antiword`, `catdoc`, `soffice`/`libreoffice`).
    - Emits log messages describing failures and fallbacks.

    Failure modes:
    - Returns `{}` when no tool is available or all tools fail/produce empty output.
    """

    last_error = None

    # antiword
    if shutil.which("antiword"):
        try:
            proc = subprocess.run(
                ["antiword", "-"],
                input=data,
                stdout=subprocess.PIPE,
                stderr=subprocess.PIPE,
                check=False,
                timeout=timeout,
            )
            text = proc.stdout.decode("utf-8", errors="ignore").strip()
            if text:
                paras = [sanitize_text(p.strip()) for p in text.split("\n\n") if p.strip()]
                return {i + 1: p for i, p in enumerate(paras)}
            else:
                last_error = "antiword returned empty output"
        except Exception as exc:  # noqa: BLE001
            last_error = f"antiword exception: {exc}"

    # catdoc
    if shutil.which("catdoc"):
        try:
            proc = subprocess.run(
                ["catdoc", "-"],
                input=data,
                stdout=subprocess.PIPE,
                stderr=subprocess.PIPE,
                check=False,
                timeout=timeout,
            )
            text = proc.stdout.decode("utf-8", errors="ignore").strip()
            if text:
                paras = [sanitize_text(p.strip()) for p in text.split("\n\n") if p.strip()]
                return {i + 1: p for i, p in enumerate(paras)}
            else:
                if last_error:
                    logger.warning("Both antiword and catdoc returned empty output for .doc")
                else:
                    last_error = "catdoc returned empty output"
        except Exception as exc:  # noqa: BLE001
            if last_error:
                logger.warning("antiword failed (%s) and catdoc exception: %s", last_error, exc)
            else:
                last_error = f"catdoc exception: {exc}"

    # soffice/libreoffice 兜底
    soffice = shutil.which("soffice") or shutil.which("libreoffice")
    if soffice:
        try:
            with tempfile.TemporaryDirectory() as tmpdir:
                src = Path(tmpdir) / "input.doc"
                out = Path(tmpdir) / "input.txt"
                src.write_bytes(data)
                proc = subprocess.run(
                    [soffice, "--headless", "--convert-to", "txt:Text", "--outdir", tmpdir, str(src)],
                    stdout=subprocess.PIPE,
                    stderr=subprocess.PIPE,
                    check=False,
                    timeout=timeout,
                )
                if out.exists():
                    text = out.read_text(encoding="utf-8", errors="ignore").strip()
                    if text:
                        paras = [sanitize_text(p.strip()) for p in text.split("\n\n") if p.strip()]
                        return {i + 1: p for i, p in enumerate(paras)}
                last_error = last_error or proc.stderr.decode("utf-8", errors="ignore") or "soffice returned no output"
        except Exception as exc:  # noqa: BLE001
            last_error = f"{last_error or ''}; soffice exception: {exc}".strip("; ")

    if last_error:
        logger.warning("Failed to extract legacy .doc: %s", last_error)
    else:
        logger.warning("No tool available to process legacy .doc (need antiword, catdoc, or soffice/libreoffice)")
    return {}


def extract_word_attachment_tasks(
    data: bytes,
    filename: str,
    content_type: str,  # 保持签名不变
    base_meta: dict,
    max_len: int,
) -> List[Tuple[str, dict]]:
    """
    Purpose:
    Produce fixed-size `(chunk_text, metadata)` tuples for a single Word attachment.

    Inputs:
    - data: Attachment bytes.
    - filename: Attachment filename (used for extension detection and logging).
    - content_type: Attachment content type (currently unused; retained for compatibility).
    - base_meta: Base metadata applied to all generated chunks.
    - max_len: Maximum chunk length in characters.

    Outputs:
    - List of `(chunk_text, metadata)` tuples; returns `[]` when unsupported or empty.

    Side effects:
    - Emits log messages describing extraction outcomes.

    Failure modes:
    - Returns `[]` when extraction fails or yields no text.
    """

    suffix = Path(filename).suffix.lower()
    if suffix not in WORD_EXTS:
        return []

    # choose extractor
    if suffix == ".docx":
        paragraphs = _extract_text_from_docx(data)
    else:  # .doc legacy
        paragraphs = extract_text_from_doc(data)

    if not paragraphs:
        return []

    # 1) 合并所有段落为一段（维持现有策略）
    full_text = join_nonempty_segments(text for _, text in sorted(paragraphs.items()))

    file_type = "docx" if suffix == ".docx" else "doc"
    tasks = build_attachment_tasks(
        full_text,
        base_meta=base_meta,
        file_type=file_type,
        filename=filename,
        max_len=max_len,
    )

    if tasks:
        logger.info("Extracted attachment %s (%d chunks)", filename, len(tasks))
    return tasks
